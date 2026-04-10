"""Tests fuer step1_analyze (JSON-Parsing, Position-Validierung, Block-ID-Mapping)."""

import json
import pytest
from src.step1_analyze import extract_json, validate_positions, resolve_block_codings
from src.pdf_extractor import (
    build_fulltext_and_positions, extract_docx, extract_document,
)
from src.run_context import RunContext


class TestExtractJson:
    def test_clean_json(self):
        text = '{"segments": [], "kernergebnisse": []}'
        result = extract_json(text)
        assert result["segments"] == []

    def test_json_in_markdown_block(self):
        text = '```json\n{"segments": [{"code_id": "A-01"}]}\n```'
        result = extract_json(text)
        assert result["segments"][0]["code_id"] == "A-01"

    def test_json_with_preamble(self):
        text = 'Hier ist die Analyse:\n\n{"segments": [], "kernergebnisse": []}'
        result = extract_json(text)
        assert "segments" in result

    def test_truncated_json_repaired(self):
        text = '{"segments": [{"code_id": "A-01", "text": "test"}, {"code_id": "B-01", "text": "halb'
        result = extract_json(text)
        assert len(result["segments"]) >= 1
        assert result["segments"][0]["code_id"] == "A-01"

    def test_no_json_raises(self):
        with pytest.raises(ValueError, match="Kein JSON"):
            extract_json("Keine JSON-Daten hier.")

    def test_top_level_array(self):
        # Vision-Classifier gibt Array zurueck
        text = '[{"page": 1, "page_type": "plan"}, {"page": 2, "page_type": "text"}]'
        result = extract_json(text)
        assert isinstance(result, list)
        assert len(result) == 2
        assert result[0]["page_type"] == "plan"

    def test_top_level_array_in_markdown(self):
        text = '```json\n[{"page": 1, "page_type": "plan"}]\n```'
        result = extract_json(text)
        assert isinstance(result, list)
        assert result[0]["page"] == 1

    def test_top_level_array_with_preamble(self):
        text = 'Hier die Klassifikation:\n[{"page": 1, "page_type": "plan"}]'
        result = extract_json(text)
        assert isinstance(result, list)
        assert len(result) == 1


class TestValidatePositions:
    def test_exact_match(self):
        full_text = "Dies ist ein Testtext mit Inhalt."
        segments = [{"text": "Testtext mit Inhalt", "char_start": 0, "char_end": 10}]
        result = validate_positions(segments, full_text)
        assert result[0]["char_start"] == 13
        assert result[0]["char_end"] == 13 + len("Testtext mit Inhalt")

    def test_partial_match_via_prefix(self):
        full_text = "Einleitung. Hier beginnt der relevante Abschnitt mit viel Text und Kontext. Ende."
        segments = [{
            "text": "Hier beginnt der relevante Abschnitt mit viel Text und Kontext. Ende. Plus extra.",
            "char_start": 0, "char_end": 50,
        }]
        result = validate_positions(segments, full_text)
        assert result[0]["char_start"] == 12

    def test_no_match_keeps_original(self):
        full_text = "Komplett anderer Text."
        segments = [{"text": "Nicht vorhanden xyz abc", "char_start": 42, "char_end": 99}]
        result = validate_positions(segments, full_text)
        assert result[0]["char_start"] == 42


class TestBlockIdMapping:
    """Tests für den Block-ID → Zeichenposition Ansatz."""

    def _make_extraction(self):
        """Erzeugt minimale Extraktionsdaten."""
        return {
            "pages": [
                {
                    "page": 1,
                    "blocks": [
                        {"id": "p1_b0", "type": "text", "text": "Erster Block",
                         "bbox": [0, 0, 100, 20]},
                        {"id": "p1_b1", "type": "text", "text": "Zweiter Block",
                         "bbox": [0, 30, 100, 50]},
                    ],
                },
                {
                    "page": 2,
                    "blocks": [
                        {"id": "p2_b0", "type": "text", "text": "Dritter Block",
                         "bbox": [0, 0, 100, 20]},
                    ],
                },
            ],
        }

    def test_build_fulltext_and_positions(self):
        data = self._make_extraction()
        fulltext, positions = build_fulltext_and_positions(data)

        # Volltext ist Blöcke verbunden mit \n\n
        assert "Erster Block" in fulltext
        assert "Zweiter Block" in fulltext
        assert "Dritter Block" in fulltext

        # Positionen stimmen
        for block_id, (start, end) in positions.items():
            block_text = fulltext[start:end]
            assert block_text.strip()  # Nicht leer

        # p1_b0 startet bei 0
        assert positions["p1_b0"][0] == 0
        assert fulltext[positions["p1_b0"][0]:positions["p1_b0"][1]] == "Erster Block"

    def test_resolve_block_codings(self):
        data = self._make_extraction()
        _, positions = build_fulltext_and_positions(data)
        block_index = {}
        for page in data["pages"]:
            for block in page["blocks"]:
                block_index[block["id"]] = block

        codings = [
            {
                "block_id": "p1_b0",
                "code_id": "A-01",
                "code_name": "Testcode",
                "hauptkategorie": "A",
            },
            {
                "block_id": "p2_b0",
                "code_id": "B-01",
                "code_name": "Anderer Code",
                "hauptkategorie": "B",
            },
        ]

        segments = resolve_block_codings(codings, positions, block_index)
        assert len(segments) == 2

        # Erstes Segment
        assert segments[0]["code_id"] == "A-01"
        assert segments[0]["text"] == "Erster Block"
        assert segments[0]["char_start"] == 0
        assert segments[0]["char_end"] == len("Erster Block")

        # Zweites Segment
        assert segments[1]["code_id"] == "B-01"
        assert segments[1]["text"] == "Dritter Block"

    def test_resolve_unknown_block_id_skipped(self):
        data = self._make_extraction()
        _, positions = build_fulltext_and_positions(data)
        block_index = {}
        for page in data["pages"]:
            for block in page["blocks"]:
                block_index[block["id"]] = block

        codings = [{"block_id": "p99_b0", "code_id": "X-01",
                     "code_name": "?", "hauptkategorie": "X"}]
        segments = resolve_block_codings(codings, positions, block_index)
        assert len(segments) == 0

    def test_multiple_codes_per_block(self):
        data = self._make_extraction()
        _, positions = build_fulltext_and_positions(data)
        block_index = {}
        for page in data["pages"]:
            for block in page["blocks"]:
                block_index[block["id"]] = block

        # Gleicher Block, zwei Codes
        codings = [
            {"block_id": "p1_b0", "code_id": "A-01",
             "code_name": "Code A", "hauptkategorie": "A"},
            {"block_id": "p1_b0", "code_id": "B-01",
             "code_name": "Code B", "hauptkategorie": "B"},
        ]
        segments = resolve_block_codings(codings, positions, block_index)
        assert len(segments) == 2
        assert segments[0]["char_start"] == segments[1]["char_start"]


class TestExtractDocx:
    """Regressionstest: .docx muss als logische Absaetze rauskommen, nicht als
    gerenderte Zeilen (pymupdf). Siehe Interview-Absatz-Regression."""

    def test_docx_paragraphs_not_lines(self, tmp_path):
        """Ein .docx mit 3 Absaetzen muss genau 3 Bloecke ergeben."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("S1: Erster Sprecher-Turn mit langem Text, "
                          "der ueber mehrere Zeilen umgebrochen wird "
                          "wenn man ihn rendered. [0:00:04.3]")
        doc.add_paragraph("S2: Zweiter Turn. [0:00:10.5]")
        doc.add_paragraph("S1: Dritter Turn. [0:00:15.0]")
        docx_path = tmp_path / "test_interview.docx"
        doc.save(str(docx_path))

        data = extract_docx(docx_path)

        blocks = data["pages"][0]["blocks"]
        assert len(blocks) == 3, (
            f"Erwartet 3 Bloecke (1 pro Absatz), bekam {len(blocks)}. "
            "Regression: .docx wird in Zeilen statt Absaetze zerlegt."
        )
        assert blocks[0]["text"].startswith("S1: Erster")
        assert blocks[1]["text"].startswith("S2: Zweiter")
        assert blocks[2]["text"].startswith("S1: Dritter")

    def test_extract_document_dispatches_docx(self, tmp_path):
        """extract_document waehlt extract_docx fuer .docx."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Hallo Welt")
        docx_path = tmp_path / "dispatch_test.docx"
        doc.save(str(docx_path))

        data = extract_document(docx_path)
        # Nur 1 synthetische Seite (kein pymupdf-Rendering)
        assert len(data["pages"]) == 1
        assert data["pages"][0]["page"] == 1

    def test_fulltext_positions_from_docx(self, tmp_path):
        """build_fulltext_and_positions funktioniert mit extract_docx-Output."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Absatz Eins")
        doc.add_paragraph("Absatz Zwei")
        docx_path = tmp_path / "pos_test.docx"
        doc.save(str(docx_path))

        data = extract_docx(docx_path)
        fulltext, positions = build_fulltext_and_positions(data)

        assert "Absatz Eins" in fulltext
        assert "Absatz Zwei" in fulltext
        assert len(positions) == 2
        for block_id, (start, end) in positions.items():
            assert fulltext[start:end].startswith("Absatz")

    def test_empty_paragraphs_skipped(self, tmp_path):
        """Leere Absaetze im .docx werden uebersprungen."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Inhalt")
        doc.add_paragraph("")  # leer
        doc.add_paragraph("   ")  # nur Whitespace
        doc.add_paragraph("Mehr Inhalt")
        docx_path = tmp_path / "empty_test.docx"
        doc.save(str(docx_path))

        data = extract_docx(docx_path)
        blocks = data["pages"][0]["blocks"]
        assert len(blocks) == 2


class TestRunContextCache:
    def test_cache_roundtrip(self, tmp_path):
        ctx = RunContext(tmp_path)
        ctx.ensure_dirs()
        data = {"segments": [{"code_id": "A-01"}], "kernergebnisse": []}

        ctx.cache_parsed("test.docx", data)
        loaded = ctx.get_cached_parsed("test.docx")

        assert loaded is not None
        assert loaded["segments"][0]["code_id"] == "A-01"

    def test_cache_miss(self, tmp_path):
        ctx = RunContext(tmp_path)
        ctx.ensure_dirs()
        assert ctx.get_cached_parsed("nonexistent.docx") is None

    def test_prompt_and_response_cached(self, tmp_path):
        ctx = RunContext(tmp_path)
        ctx.ensure_dirs()

        ctx.cache_prompt("test.docx", "Der Prompt")
        ctx.cache_response("test.docx", "Die Antwort")

        assert (ctx.prompts_dir / "test.docx.txt").exists()
        assert (ctx.responses_dir / "test.docx.txt").exists()

    def test_state_tracking(self, tmp_path):
        ctx = RunContext(tmp_path)
        ctx.ensure_dirs()
        ctx.init_state("mayring", None, ["a.docx", "b.docx"])

        assert ctx.get_pending_transcripts() == ["a.docx", "b.docx"]

        ctx.mark_transcript_done("a.docx")
        assert ctx.get_pending_transcripts() == ["b.docx"]

        ctx.mark_step_done(1)
        assert ctx.is_step_done(1)
        assert not ctx.is_step_done(2)

    def test_interrupted_detection(self, tmp_path, monkeypatch):
        from src import run_context
        monkeypatch.setattr(run_context, "OUTPUT_ROOT", tmp_path)

        # Erstelle einen "unterbrochenen" run
        run_dir = tmp_path / "2026-01-01_12-00-00"
        ctx = RunContext(run_dir)
        ctx.ensure_dirs()
        ctx.init_state("mayring", None, ["test.docx"])

        from src.run_context import find_interrupted_runs
        interrupted = find_interrupted_runs()
        assert len(interrupted) == 1
        assert interrupted[0].run_dir == run_dir

        # Nach complete: kein Interrupted mehr
        ctx.mark_completed()
        interrupted = find_interrupted_runs()
        assert len(interrupted) == 0
