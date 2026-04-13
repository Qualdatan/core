"""Tests fuer den Office-Konverter.

Die Backend-spezifischen Tests werden uebersprungen, wenn weder
PowerShell-COM (Word/Excel) noch LibreOffice verfuegbar ist.
"""

from pathlib import Path

import pytest

from qualdatan_core.office_converter import (
    SUPPORTED_EXTENSIONS,
    OfficeConverterUnavailable,
    convert_to_pdf,
    detect_backend,
    find_office_files,
    reset_backend_cache,
)


@pytest.fixture(autouse=True)
def _reset_cache():
    reset_backend_cache()
    yield
    reset_backend_cache()


@pytest.fixture
def docx_file(tmp_path):
    """Erzeugt eine kleine .docx via python-docx."""
    pytest.importorskip("docx")
    from docx import Document

    src = tmp_path / "test.docx"
    doc = Document()
    doc.add_heading("Testdokument", 0)
    doc.add_paragraph("Erste Zeile")
    doc.add_paragraph("Zweite Zeile mit Umlauten: aeoeue")
    doc.save(str(src))
    return src


@pytest.fixture
def xlsx_multi_sheet(tmp_path):
    """Erzeugt eine .xlsx mit drei Sheets via openpyxl."""
    pytest.importorskip("openpyxl")
    import openpyxl

    src = tmp_path / "multi.xlsx"
    wb = openpyxl.Workbook()
    s1 = wb.active
    s1.title = "Sheet1"
    s1.append(["A", "B", "C"])
    for i in range(20):
        s1.append([i, i * 2, f"row {i}"])
    s2 = wb.create_sheet("Sheet2")
    s2.append(["LPH", "Stunden"])
    for i in range(1, 8):
        s2.append([f"LPH {i}", i * 100])
    s3 = wb.create_sheet("Notizen")
    s3.append(["Eine Anmerkung"])
    wb.save(str(src))
    return src


def _backend_available() -> bool:
    return detect_backend() is not None


# ---------------------------------------------------------------------------
# Backend-Erkennung (immer ausfuehrbar)
# ---------------------------------------------------------------------------

class TestDetectBackend:
    def test_detect_returns_known_backend_or_none(self):
        result = detect_backend()
        assert result in (None, "powershell-com", "libreoffice")

    def test_detect_is_cached(self):
        first = detect_backend()
        second = detect_backend()
        assert first == second


# ---------------------------------------------------------------------------
# find_office_files
# ---------------------------------------------------------------------------

class TestFindOfficeFiles:
    def test_finds_supported_extensions(self, tmp_path):
        (tmp_path / "a.docx").write_bytes(b"")
        (tmp_path / "b.xlsx").write_bytes(b"")
        (tmp_path / "c.pdf").write_bytes(b"")  # nicht Office
        (tmp_path / "d.txt").write_bytes(b"")  # nicht Office
        results = find_office_files(tmp_path)
        names = [p.name for p in results]
        assert "a.docx" in names
        assert "b.xlsx" in names
        assert "c.pdf" not in names

    def test_ignores_lock_files(self, tmp_path):
        (tmp_path / "real.docx").write_bytes(b"")
        (tmp_path / "~$real.docx").write_bytes(b"")
        results = find_office_files(tmp_path)
        names = [p.name for p in results]
        assert "real.docx" in names
        assert "~$real.docx" not in names

    def test_recursive(self, tmp_path):
        sub = tmp_path / "deep" / "nested"
        sub.mkdir(parents=True)
        (sub / "x.xlsx").write_bytes(b"")
        results = find_office_files(tmp_path)
        assert len(results) == 1
        assert results[0].name == "x.xlsx"

    def test_project_filter(self, tmp_path):
        (tmp_path / "A").mkdir()
        (tmp_path / "A" / "x.docx").write_bytes(b"")
        (tmp_path / "B").mkdir()
        (tmp_path / "B" / "y.docx").write_bytes(b"")
        results = find_office_files(tmp_path, project_filter="A")
        assert len(results) == 1
        assert results[0].name == "x.docx"

    def test_empty_directory(self, tmp_path):
        assert find_office_files(tmp_path) == []

    def test_supported_extensions_list(self):
        assert ".docx" in SUPPORTED_EXTENSIONS
        assert ".xlsx" in SUPPORTED_EXTENSIONS
        assert ".doc" in SUPPORTED_EXTENSIONS
        assert ".xls" in SUPPORTED_EXTENSIONS


# ---------------------------------------------------------------------------
# convert_to_pdf — fehleresillient ohne Backend
# ---------------------------------------------------------------------------

class TestConvertToPdfErrorHandling:
    def test_unsupported_extension_raises(self, tmp_path):
        src = tmp_path / "test.txt"
        src.write_bytes(b"hello")
        with pytest.raises(ValueError, match="Nicht unterstuetzte Extension"):
            convert_to_pdf(src, tmp_path / "out.pdf")

    def test_missing_source_raises(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            convert_to_pdf(tmp_path / "missing.docx", tmp_path / "out.pdf")

    def test_no_backend_raises(self, tmp_path, monkeypatch):
        monkeypatch.setattr(
            "src.office_converter.detect_backend", lambda: None
        )
        src = tmp_path / "x.docx"
        src.write_bytes(b"placeholder")
        with pytest.raises(OfficeConverterUnavailable):
            convert_to_pdf(src, tmp_path / "out.pdf")


# ---------------------------------------------------------------------------
# Integrationstests: nur wenn Backend verfuegbar
# ---------------------------------------------------------------------------

@pytest.mark.skipif(not _backend_available(),
                    reason="Kein Office-Backend installiert")
class TestRealConversion:
    def test_docx_to_pdf(self, docx_file, tmp_path):
        dst = tmp_path / "out.pdf"
        result = convert_to_pdf(docx_file, dst)
        assert result == dst
        assert dst.exists()
        assert dst.stat().st_size > 1000  # nicht leer

        # Lesbar als PDF?
        import fitz
        doc = fitz.open(str(dst))
        assert len(doc) >= 1
        text = doc[0].get_text()
        assert "Testdokument" in text
        doc.close()

    def test_xlsx_multi_sheet_produces_multiple_pages(self, xlsx_multi_sheet, tmp_path):
        dst = tmp_path / "multi.pdf"
        convert_to_pdf(xlsx_multi_sheet, dst)
        assert dst.exists()

        import fitz
        doc = fitz.open(str(dst))
        # 3 Sheets → mindestens 3 Seiten
        assert len(doc) >= 3
        all_text = "\n".join(p.get_text() for p in doc)
        assert "LPH" in all_text  # Sheet 2 Inhalt
        assert "Notizen" in all_text or "Anmerkung" in all_text  # Sheet 3
        doc.close()

    def test_cache_skips_when_dst_newer(self, docx_file, tmp_path):
        dst = tmp_path / "cached.pdf"
        convert_to_pdf(docx_file, dst)
        first_mtime = dst.stat().st_mtime

        # Zweiter Aufruf sollte Cache benutzen (mtime bleibt)
        convert_to_pdf(docx_file, dst)
        assert dst.stat().st_mtime == first_mtime

    def test_force_overrides_cache(self, docx_file, tmp_path):
        dst = tmp_path / "forced.pdf"
        convert_to_pdf(docx_file, dst)
        first_mtime = dst.stat().st_mtime

        import time
        time.sleep(1.1)  # Filesystem-Aufloesung
        convert_to_pdf(docx_file, dst, force=True)
        assert dst.stat().st_mtime > first_mtime
